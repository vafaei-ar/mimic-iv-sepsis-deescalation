#!/usr/bin/env python3
"""Build an Intensive Care Medicine original-paper submission package.

Consumes only frozen aggregate publication outputs. No patient-level data are read, no models
are refit, and no estimands are redefined. The package follows current ICM original-paper
constraints: <=3000 main-text words, <=5 tables/figures, <=40 references, structured abstract,
4-6 keywords, two-sentence take-home message, and a <=140-character social-media summary.
"""
from __future__ import annotations

from pathlib import Path
import re

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HARM = Path("outputs/publication_integration/harmonized")
OUT = Path("outputs/publication_integration/icm_package")


def pct(x: float) -> float:
    return 100.0 * float(x)


def save_fig(fig: plt.Figure, stem: str) -> None:
    fig.savefig(OUT / f"{stem}.png", dpi=300, bbox_inches="tight")
    fig.savefig(OUT / f"{stem}.pdf", bbox_inches="tight")
    plt.close(fig)


def word_count(text: str) -> int:
    text = re.sub(r"[#*`|]", " ", text)
    return len(re.findall(r"\b[\w’'-]+\b", text))


def make_timeline() -> None:
    fig, ax = plt.subplots(figsize=(7.2, 2.1))
    ax.set_xlim(-4, 124)
    ax.set_ylim(-0.7, 1.0)
    ax.axis("off")
    ax.hlines(0, 0, 120, linewidth=1.5, color="#1f4e79")
    ax.annotate("", xy=(120, 0), xytext=(116, 0), arrowprops=dict(arrowstyle="->", lw=1.5, color="#1f4e79"))
    for x, label in [
        (0, "First qualifying\nbroad-spectrum antibiotic"),
        (72, "Day-3\ndecision"),
        (96, "Landmark /\nfollow-up start"),
        (120, "30-day outcome\nhorizon continues"),
    ]:
        ax.vlines(x, -0.12, 0.12, linewidth=1.5, color="#1f4e79")
        ax.text(x, 0.18, label, ha="center", va="bottom", fontsize=8)
    ax.plot([0, 72], [-0.18, -0.18], lw=4, color="#9ecae1", solid_capstyle="butt")
    ax.plot([72, 96], [-0.18, -0.18], lw=4, color="#4f81bd", solid_capstyle="butt")
    ax.plot([96, 120], [-0.18, -0.18], lw=4, color="#1f4e79", solid_capstyle="butt")
    ax.text(36, -0.36, "Pre-decision covariates", ha="center", fontsize=8)
    ax.text(84, -0.36, "Exposure classification (72–96 h)", ha="center", fontsize=8)
    ax.text(108, -0.36, "Outcome follow-up", ha="center", fontsize=8)
    fig.tight_layout()
    save_fig(fig, "Fig1_target_trial_timeline")


def make_progressive(progressive: pd.DataFrame) -> None:
    d = progressive.copy()
    y = np.arange(len(d))[::-1]
    est = 100 * d["risk_difference"].to_numpy(float)
    lo = 100 * d["rd_lower_95"].to_numpy(float)
    hi = 100 * d["rd_upper_95"].to_numpy(float)
    xerr = np.vstack([est - lo, hi - est])
    labels = [
        "M1 demographics/comorbidity",
        "M2 + baseline severity/labs",
        "M3 + near-decision status",
        "M4 + trajectories/intensity",
    ]
    fig, ax = plt.subplots(figsize=(7.2, 3.4))
    ax.errorbar(est, y, xerr=xerr, fmt="o", capsize=3, lw=1.3, color="#1f4e79", ecolor="#4f81bd")
    ax.axvline(0, lw=1, linestyle="--", color="#666666")
    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_xlabel("30-day mortality risk difference, percentage points", fontsize=9)
    ax.grid(axis="x", linewidth=0.4, alpha=0.3)
    for yi, e in zip(y, est):
        ax.text(e + 0.18, yi, f"{e:+.2f}", va="center", fontsize=8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    save_fig(fig, "Fig2_progressive_adjustment")


def make_cross_dataset(mortality: pd.DataFrame, secondary: pd.DataFrame) -> None:
    fig, axes = plt.subplots(1, 2, figsize=(8.5, 3.6))
    ax = axes[0]
    d = mortality.iloc[:2]
    est = 100 * d["mortality_rd"].to_numpy(float)
    lo = 100 * d["rd_ci95_low"].to_numpy(float)
    hi = 100 * d["rd_ci95_high"].to_numpy(float)
    y = np.array([1, 0])
    ax.errorbar(est, y, xerr=np.vstack([est-lo, hi-est]), fmt="o", capsize=3, lw=1.3,
                color="#1f4e79", ecolor="#4f81bd")
    ax.axvline(0, lw=1, linestyle="--", color="#666666")
    ax.set_yticks(y)
    ax.set_yticklabels(["MIMIC-IV", "Penn State"], fontsize=8)
    ax.set_xlabel("Mortality RD, percentage points", fontsize=8)
    ax.set_title("a  30-day mortality", loc="left", fontsize=9, fontweight="bold")
    ax.grid(axis="x", linewidth=0.4, alpha=0.3)

    ax = axes[1]
    names = ["Antibiotic-free days", "Normalized systemic antibiotic exposure", "Normalized broad-spectrum exposure"]
    labels = ["Antibiotic-free days", "Systemic exposure", "Broad-spectrum exposure"]
    # Standardize only for compact visual direction: plot raw estimates in separate rows, annotated with units.
    ys = np.arange(6)[::-1]
    positions, vals, los, his, txt = [], [], [], [], []
    k = 0
    for name, label in zip(names, labels):
        dd = secondary[secondary["outcome"] == name]
        for _, r in dd.iterrows():
            positions.append(ys[k]); vals.append(float(r.estimate)); los.append(float(r.ci95_low)); his.append(float(r.ci95_high))
            txt.append(f"{r.dataset}: {label}"); k += 1
    vals = np.array(vals); los = np.array(los); his = np.array(his)
    # Plot antibiotic-free days on top x axis and normalized exposure on bottom x axis would confuse; instead annotate estimates.
    ax.axis("off")
    ax.set_title("b  Stewardship outcomes", loc="left", fontsize=9, fontweight="bold")
    y0 = 0.87
    for i, (name, label) in enumerate(zip(names, labels)):
        dd = secondary[secondary["outcome"] == name]
        ax.text(0.02, y0, label, transform=ax.transAxes, fontsize=8.5, fontweight="bold", va="top")
        y0 -= 0.10
        for _, r in dd.iterrows():
            unit = "days" if r.unit == "days" else ""
            ax.text(0.05, y0, f"{r.dataset}: {r.estimate:+.2f} ({r.ci95_low:+.2f}, {r.ci95_high:+.2f}) {unit}".rstrip(),
                    transform=ax.transAxes, fontsize=8, va="top")
            y0 -= 0.085
        y0 -= 0.04
    fig.tight_layout()
    save_fig(fig, "Fig3_cross_dataset_results")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)


def build_docx(title: str, short_title: str, abstract_sections: list[tuple[str,str]], keywords: list[str],
               take_home: str, tweet: str, body_sections: list[tuple[str,str]], refs: list[str],
               table1_rows: list[list[str]], table2_rows: list[list[str]]) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8); sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"; styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Times New Roman"; styles["Title"].font.size = Pt(14)
    for s in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[s].font.name = "Times New Roman"
    p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(title)
    doc.add_paragraph("Author names: Alireza Vafaei Sadr, [student collaborators], Anthony S. Bonavia, Vida Abedi, and [additional collaborators]")
    doc.add_paragraph("Affiliations: Department of Public Health Sciences, College of Medicine, Pennsylvania State University, Hershey, Pennsylvania, USA; [additional affiliations]")
    doc.add_paragraph("Corresponding author: Alireza Vafaei Sadr, [email], [telephone]")
    doc.add_paragraph(f"Running title: {short_title}")
    doc.add_paragraph("Conflicts of interest: [Insert final statement before submission]")
    doc.add_heading("Abstract", level=1)
    for label, text in abstract_sections:
        p = doc.add_paragraph(); r = p.add_run(label + ": "); r.bold = True; p.add_run(text)
    p = doc.add_paragraph(); r = p.add_run("Keywords: "); r.bold = True; p.add_run("; ".join(keywords))
    doc.add_heading("Take-home message", level=1); doc.add_paragraph(take_home)
    doc.add_heading("140-character social-media summary", level=1); doc.add_paragraph(tweet)
    for heading, text in body_sections:
        doc.add_heading(heading, level=1)
        for para in text.split("\n\n"):
            if para.strip(): doc.add_paragraph(para.strip())
    doc.add_heading("Table 1. Target trial emulation and modified external replication", level=1)
    add_table(doc, ["Protocol element", "MIMIC-IV", "Penn State modified replication"], table1_rows, [1.4, 2.9, 2.9])
    doc.add_heading("Table 2. Primary and secondary outcomes", level=1)
    add_table(doc, ["Outcome", "MIMIC-IV", "Penn State"], table2_rows, [2.7, 2.2, 2.2])
    doc.add_heading("Figure legends", level=1)
    doc.add_paragraph("Fig. 1 Target-trial timing and landmark design. Covariates were measured before the 72-hour decision; treatment was classified during 72–96 hours; follow-up began at 96 hours.")
    doc.add_paragraph("Fig. 2 Progressive adjustment of the MIMIC-IV mortality association. Risk differences are de-escalation/stopping minus continued broad-spectrum therapy. Error bars show 95% confidence intervals.")
    doc.add_paragraph("Fig. 3 Cross-dataset results. a, Primary 30-day mortality risk differences. b, Stewardship outcomes. Penn State is a modified external replication and estimates are not pooled.")
    doc.add_heading("References", level=1)
    for i, ref in enumerate(refs, 1): doc.add_paragraph(f"{i}. {ref}")
    doc.add_page_break()
    doc.add_heading("Electronic supplementary material plan", level=1)
    for item in [
        "ESM Table 1. Detailed cohort construction and exposure groups.",
        "ESM Table 2. Progressive adjustment with risk ratios, confidence intervals and balance diagnostics.",
        "ESM Table 3. MIMIC weighting sensitivities and Penn State MED_ADMIN/lenient-landmark robustness analyses.",
        "ESM Table 4. Antibiotic definitions and medication exposure audit.",
        "ESM Table 5. Microbiology algorithm and strict test-name culture sensitivity audit.",
        "ESM Table 6. Covariate timing, missingness and balance diagnostics.",
        "ESM Table 7. TARGET and RECORD reporting checklists.",
    ]: doc.add_paragraph(item)
    doc.save(OUT / "ICM_manuscript_draft.docx")


def main() -> None:
    required = [HARM / "harmonized_mortality_results.csv", HARM / "harmonized_secondary_outcomes.csv",
                HARM / "mimic_progressive_adjustment.csv", HARM / "weighting_diagnostics.csv"]
    missing = [str(p) for p in required if not p.exists()]
    if missing: raise SystemExit("Missing frozen publication inputs: " + ", ".join(missing))
    OUT.mkdir(parents=True, exist_ok=True)
    mortality = pd.read_csv(required[0]); secondary = pd.read_csv(required[1]); progressive = pd.read_csv(required[2]); weights = pd.read_csv(required[3])
    m, p = mortality.iloc[0], mortality.iloc[1]
    mw = weights[weights.dataset == "MIMIC-IV"].iloc[0]; pw = weights[weights.dataset == "PSU"].iloc[0]
    sec = {(r.dataset, r.outcome): r for r in secondary.itertuples(index=False)}

    title = "Day-3 antibiotic de-escalation in suspected ICU infection: a target trial emulation"
    short_title = "Day-3 antibiotic de-escalation"
    keywords = ["antimicrobial stewardship", "sepsis", "intensive care", "target trial emulation", "de-escalation", "MIMIC-IV"]
    abstract_sections = [
        ("Purpose", "To estimate mortality and antibiotic-burden outcomes after a day-3 broad-spectrum antibiotic de-escalation strategy while accounting for pre-decision clinical improvement."),
        ("Methods", "We emulated a 72-hour treatment decision with a 96-hour landmark in MIMIC-IV and performed a modified external replication in Penn State clinical data. Stabilized inverse probability treatment weighting adjusted for demographics, comorbidity, physiology, laboratory values, clinical trajectories and treatment intensity measured before 72 hours. The primary outcome was 30-day mortality."),
        ("Results", f"MIMIC-IV included 9,589 admissions (1,863 de-escalated/stopped; 7,726 continued). Fully adjusted mortality was 18.3% versus 17.5% (risk difference {pct(m.mortality_rd):+.2f} percentage points, 95% CI {pct(m.rd_ci95_low):+.2f} to {pct(m.rd_ci95_high):+.2f}; risk ratio {m.mortality_rr:.3f}). The apparent mortality advantage attenuated from -4.14 to +0.84 percentage points with richer adjustment. In Penn State (n=19,841), the modified-replication mortality risk difference was {pct(p.mortality_rd):+.2f} percentage points (95% CI {pct(p.rd_ci95_low):+.2f} to {pct(p.rd_ci95_high):+.2f}). Antibiotic-free days increased and systemic and broad-spectrum exposure decreased in both datasets."),
        ("Conclusions", "After adjustment for clinical trajectory and treatment intensity, MIMIC-IV showed no clear mortality benefit or harm with day-3 de-escalation. The Penn State estimate was more favorable but is not directly exchangeable because eligibility and measurement differed. Lower antibiotic burden was the most consistent cross-dataset finding."),
    ]
    abstract_text = " ".join(t for _, t in abstract_sections)
    take_home = ("In MIMIC-IV, the apparent survival advantage of day-3 antibiotic de-escalation disappeared after adjustment for pre-decision clinical trajectory and treatment intensity. "
                 "Across MIMIC-IV and a modified Penn State replication, de-escalation consistently reduced subsequent antibiotic burden, while mortality estimates differed across data representations.")
    tweet = "Day-3 ICU antibiotic de-escalation lowered antibiotic burden; MIMIC mortality benefit disappeared after trajectory adjustment."

    intro = ("Early broad-spectrum antibiotics are central to suspected sepsis care, but continued exposure after the first several days increases treatment burden and selection pressure. International guidance therefore supports reassessment and de-escalation when clinical and microbiological information becomes available [1, 2]. Culture-negative suspected ICU infection is a particularly difficult setting because negative cultures reduce microbiological justification for broad therapy without proving absence of infection.\n\n"
             "Observational evidence on de-escalation is difficult to interpret because the treatment decision is linked to recovery. Patients whose physiology and organ dysfunction improve before day 3 are more likely to have antibiotics narrowed or stopped and are also more likely to survive. Prior studies generally have not shown increased mortality after de-escalation, but exposure definitions and adjustment for evolving clinical status vary [3-6]. Recent cohorts further emphasize the importance of early clinical trajectories [12, 13].\n\n"
             "We therefore emulated a day-3 treatment decision in MIMIC-IV using a 96-hour landmark and explicit adjustment for pre-decision clinical trajectories and treatment intensity. We evaluated transportability in a modified external replication using Penn State clinical data, preserving the decision and landmark logic where feasible while documenting measurement differences that prevented exact replication.")

    methods = ("We conducted a retrospective target trial emulation in MIMIC-IV version 3.1 and a modified external replication in Penn State clinical data. In MIMIC-IV, the trial clock began at first qualifying prescribed systemic intravenous broad-spectrum antibiotic exposure. Eligible adult ICU admissions had clinical microbiology sampled, no positive clinical culture result available by 72 hours, broad-spectrum coverage during 48–72 hours, no active vasopressor overlap during the preceding 6 hours, and remained alive and hospitalized through the 96-hour landmark. Treatment was classified during 72–96 hours as no broad-spectrum overlap (de-escalation/stopping) versus any overlap (continuation). Follow-up began at 96 hours (Table 1; Fig. 1).\n\n"
               "All propensity-score covariates preceded the 72-hour decision and included demographics, comorbidity, baseline and near-decision physiology and laboratory values, organ-dysfunction and vital-sign trajectories, vasopressor trajectories, diagnostic and microbiology intensity, antibiotic intensity, steroids and BMI proxy where available. Stabilized inverse probability treatment weights estimated the average treatment effect. Balance was assessed by absolute standardized mean differences, effective sample size and maximum weight. Percentile bootstrap confidence intervals used 1,000 resamples. The primary outcome was all-cause mortality within 30 days after the landmark; secondary outcomes included hospital-free days, antibiotic-free days and normalized systemic and broad-spectrum antibiotic exposure.\n\n"
               "Penn State retained the 72-hour decision, 72–96-hour exposure window and 96-hour landmark but used a systemic broad-spectrum PRESCRIBING proxy, hospital-level timing and no faithful day-3 culture-result-availability restriction. MED_ADMIN exposure and a lenient landmark were prespecified sensitivities. The external analysis is therefore a modified replication rather than an exact transport of the MIMIC estimand. Full phenotype definitions, weighting diagnostics, sensitivity analyses and TARGET/RECORD reporting details are provided in the electronic supplementary material.")

    m_af = sec[("MIMIC-IV", "Antibiotic-free days")]; p_af = sec[("PSU", "Antibiotic-free days")]
    m_sys = sec[("MIMIC-IV", "Normalized systemic antibiotic exposure")]; p_sys = sec[("PSU", "Normalized systemic antibiotic exposure")]
    m_broad = sec[("MIMIC-IV", "Normalized broad-spectrum exposure")]; p_broad = sec[("PSU", "Normalized broad-spectrum exposure")]
    results = (f"The corrected MIMIC-IV cohort included 9,589 admissions: 1,863 (19.4%) de-escalated/stopped and 7,726 continued broad-spectrum therapy. The Penn State strict modified-replication cohort included 19,841 encounters, with 5,346 de-escalated and 14,495 continued.\n\n"
               f"In MIMIC-IV, the mortality association attenuated as richer pre-decision information was added: the 30-day mortality risk difference moved from -4.14 percentage points (95% CI -5.97 to -2.29) with demographics/comorbidity to -3.19 after baseline severity, -2.01 after near-decision clinical status, and +0.84 after improvement trajectories and treatment intensity (Fig. 2). In the corrected fully adjusted model, weighted 30-day mortality was 18.3% under de-escalation/stopping and 17.5% under continuation, risk difference {pct(m.mortality_rd):+.2f} percentage points (95% CI {pct(m.rd_ci95_low):+.2f} to {pct(m.rd_ci95_high):+.2f}) and risk ratio {m.mortality_rr:.3f}. The confidence interval was compatible with benefit and harm and does not establish equivalence. Maximum post-weighting absolute SMD was {mw.max_post_smd:.3f}, treated effective sample size {mw.ess_deescalated:.0f}, and maximum weight {mw.max_weight:.2f}.\n\n"
               f"The Penn State primary modified-replication mortality risk difference was {pct(p.mortality_rd):+.2f} percentage points (95% CI {pct(p.rd_ci95_low):+.2f} to {pct(p.rd_ci95_high):+.2f}) and risk ratio {p.mortality_rr:.3f} (95% CI {p.rr_ci95_low:.3f} to {p.rr_ci95_high:.3f}). MED_ADMIN and lenient-landmark sensitivities were similar. Measured balance was stronger than in MIMIC-IV (maximum post-weighting absolute SMD {pw.max_post_smd:.3f}), but the populations and measurement processes were not equivalent (Fig. 3).\n\n"
               f"Antibiotic burden decreased consistently. Antibiotic-free days increased by {m_af.estimate:.2f} days (95% CI {m_af.ci95_low:.2f} to {m_af.ci95_high:.2f}) in MIMIC-IV and {p_af.estimate:.2f} days (95% CI {p_af.ci95_low:.2f} to {p_af.ci95_high:.2f}) in Penn State. Normalized systemic exposure decreased by {abs(m_sys.estimate):.3f} and {abs(p_sys.estimate):.3f}, respectively, and broad-spectrum exposure by {abs(m_broad.estimate):.3f} and {abs(p_broad.estimate):.3f} (Table 2; Fig. 3).")

    discussion = ("The central finding is not that de-escalation improves survival. In MIMIC-IV, an apparent mortality advantage under limited adjustment progressively disappeared after clinical improvement, near-decision status and treatment intensity were represented. This attenuation is consistent with substantial confounding by indication and recovery trajectory. The fully adjusted estimate therefore provides no clear evidence of mortality benefit or harm and should not be interpreted as proof of safety or equivalence. Residual balance and positivity limitations reinforce this caution.\n\n"
                  "The Penn State replication produced a more favorable mortality estimate, but the two analyses should not be forced into numerical agreement. Penn State lacked a faithful day-3 positive-culture availability rule, exact ICU timing and verified intravenous route, and several outcomes relied on calendar-date semantics. These differences change the eligible population and measurement process. Strong measured balance does not remove unmeasured confounding, so the persistent lower-mortality association remains a robust observational finding within that representation rather than evidence that de-escalation causally reduces mortality.\n\n"
                  "The most reproducible cross-dataset signal was stewardship-related. De-escalation was associated with more antibiotic-free days and lower systemic and broad-spectrum exposure in both datasets. These outcomes are mechanically downstream of the treatment strategy and should not be interpreted as independent proof of clinical benefit; rather, they quantify treatment separation and the stewardship consequence of the day-3 decision.\n\n"
                  "Our study adds to recent culture-negative sepsis and ICU cohorts by explicitly demonstrating how the mortality association changes across progressively richer pre-decision adjustment [12, 13]. Strengths include a target-trial and landmark framework, confounder measurement restricted to the pre-decision period, direct representation of clinical trajectories, 1,000-replicate bootstrap inference and a frozen external replication with prespecified robustness checks. Limitations include prescription-based MIMIC exposure, selection through the 96-hour landmark, residual imbalance and low treated effective sample size, incomplete measurement of clinician judgment and source control, and substantial Penn State measurement differences. Prospective or more closely harmonized multisite evaluation is needed for stronger causal conclusions about mortality.")

    conclusion = "After adjustment for pre-decision clinical improvement and treatment intensity, day-3 de-escalation in MIMIC-IV was not clearly associated with higher or lower 30-day mortality, while antibiotic burden was lower. A modified Penn State replication showed a lower-mortality association, but measurement and eligibility differences preclude interpreting this as confirmation of a causal survival benefit."
    availability = ("MIMIC-IV is available to credentialed users through PhysioNet under its data-use agreement. Penn State source data are not publicly available because of institutional data-use restrictions. Analysis code, phenotype definitions, aggregate tables and figure-generation code will be available in the public project repository before submission. [Insert repository URL.]\n\n"
                    "Ethics: MIMIC-IV contains de-identified data. The Penn State modified external replication used institutional clinical data under local governance. [Insert applicable Penn State IRB protocol number or determination.]\n\n"
                    "Funding: [Insert final funding statement.]\n\nCompeting interests: [Insert final competing-interest statement.]\n\nAcknowledgments: The authors acknowledge the developers and maintainers of MIMIC-IV and PhysioNet. [Add collaborators and institutional acknowledgments as appropriate.]")

    refs = [
        "Evans L, Rhodes A, Alhazzani W et al (2021) Surviving sepsis campaign: international guidelines for management of sepsis and septic shock 2021. Intensive Care Med 47:1181–1247",
        "Tabah A, Bassetti M, Kollef MH et al (2020) Antimicrobial de-escalation in critically ill patients: a position statement from the ESICM and ESCMID Critically Ill Patients Study Group. Intensive Care Med 46:245–265",
        "Guo Y, Gao W, Yang H, Ma C, Sui S (2016) De-escalation of empiric antibiotics in patients with severe sepsis or septic shock: a meta-analysis. Heart Lung 45:454–459",
        "Roper S, Wingler MJB, Cretella DA (2023) Antibiotic de-escalation in critically ill patients with negative clinical cultures. Pharmacy 11:104",
        "Kim YC, Kim JH, Ahn JY et al (2020) Discontinuation of glycopeptides in patients with culture negative severe sepsis or septic shock: a propensity-matched retrospective cohort study. Antibiotics 9:250",
        "Patanwala AE, Abu Sardaneh A, Alffenaar J-WC et al (2025) Antibiotic de-escalation practices in the intensive care unit: a multicenter observational study. Ann Pharmacother 59:311–318",
        "Johnson AEW, Bulgarelli L, Shen L et al (2023) MIMIC-IV, a freely accessible electronic health record dataset. Sci Data 10:1",
        "Matthews AA, Danaei G, Islam N, Kurth T (2022) Target trial emulation: applying principles of randomised trials to observational studies. BMJ 378:e071108",
        "Cashin AG, Hansford HJ, Hernán MA et al (2025) Transparent reporting of observational studies emulating a target trial—the TARGET Statement. JAMA 334:1084–1093",
        "Benchimol EI, Smeeth L, Guttmann A et al (2015) The REporting of studies Conducted using Observational Routinely collected health Data (RECORD) Statement. PLoS Med 12:e1001885",
        "Burrows P, Brown R-A, Samuelsen A, Bonavia AS (2025) Association between in-hospital antibiotic use and long-term outcomes in critically ill patients. Antimicrob Steward Healthc Epidemiol. doi:10.1017/ash.2025.10054",
        "Ohnuma T, Fuller M, Balamurugan P et al (2026) Antibiotic de-escalation and 30-day mortality in patients with suspected bacterial culture-negative sepsis. J Crit Care 95:155646",
        "Makino J, Honda M, Sato F et al (2026) Clinical impact of antimicrobial de-escalation in critically ill patients: a single-center cohort study. J Infect Chemother 32:102946",
    ]

    table1 = [
        ["Eligibility", "Adult ICU; early prescribed systemic IV broad-spectrum therapy; microbiology sampled; no positive culture result available by 72 h; stable decision eligibility; alive/hospitalized through 96 h", "Same decision/landmark concepts where feasible; no faithful day-3 culture-result rule; hospital-level timing"],
        ["Strategies", "No broad-spectrum overlap during 72–96 h vs any overlap", "Primary PRESCRIBING broad-spectrum proxy; MED_ADMIN sensitivity"],
        ["Assignment", "72 h after first qualifying broad-spectrum exposure", "Conceptual 72-h decision"],
        ["Follow-up", "From 96-h landmark", "From approximated 96-h landmark"],
        ["Primary outcome", "30-day all-cause mortality", "30-day all-cause mortality"],
        ["Analysis", "Stabilized IPTW ATE; 1,000-replicate bootstrap", "Frozen harmonized IPTW ATE; 1,000-replicate bootstrap"],
    ]
    table2 = [
        ["30-day mortality RD", f"{pct(m.mortality_rd):+.2f} pp ({pct(m.rd_ci95_low):+.2f}, {pct(m.rd_ci95_high):+.2f})", f"{pct(p.mortality_rd):+.2f} pp ({pct(p.rd_ci95_low):+.2f}, {pct(p.rd_ci95_high):+.2f})"],
        ["30-day mortality RR", f"{m.mortality_rr:.3f}", f"{p.mortality_rr:.3f} ({p.rr_ci95_low:.3f}, {p.rr_ci95_high:.3f})"],
        ["Hospital-free days", f"{sec[(('MIMIC-IV','Hospital-free days'))].estimate:+.2f}", f"{sec[(('PSU','Hospital-free days'))].estimate:+.2f}"],
        ["Antibiotic-free days", f"{m_af.estimate:+.2f} ({m_af.ci95_low:+.2f}, {m_af.ci95_high:+.2f})", f"{p_af.estimate:+.2f} ({p_af.ci95_low:+.2f}, {p_af.ci95_high:+.2f})"],
        ["Normalized systemic exposure", f"{m_sys.estimate:+.3f}", f"{p_sys.estimate:+.3f}"],
        ["Normalized broad-spectrum exposure", f"{m_broad.estimate:+.3f}", f"{p_broad.estimate:+.3f}"],
    ]

    body_sections = [("Introduction", intro), ("Methods", methods), ("Results", results), ("Discussion", discussion), ("Conclusions", conclusion), ("Declarations", availability)]
    main_text = "\n\n".join(t for h,t in body_sections if h in {"Introduction","Methods","Results","Discussion","Conclusions"})
    if not 150 <= word_count(abstract_text) <= 250: raise SystemExit(f"Structured abstract word count out of range: {word_count(abstract_text)}")
    if word_count(main_text) > 3000: raise SystemExit(f"Main text exceeds ICM 3000-word limit: {word_count(main_text)}")
    if len(refs) > 40: raise SystemExit("References exceed ICM limit")
    if len(tweet) > 140: raise SystemExit(f"Tweet exceeds 140 characters: {len(tweet)}")

    md = f"# {title}\n\n**Running title:** {short_title}\n\n**Authors:** Alireza Vafaei Sadr, [student collaborators], Anthony S. Bonavia, Vida Abedi, and [additional collaborators]\n\n**Affiliations:** Department of Public Health Sciences, College of Medicine, Pennsylvania State University, Hershey, Pennsylvania, USA; [additional affiliations]\n\n**Corresponding author:** Alireza Vafaei Sadr, [email], [telephone]\n\n## Abstract\n\n" + "\n\n".join(f"**{k}:** {v}" for k,v in abstract_sections) + f"\n\n**Keywords:** {'; '.join(keywords)}\n\n## Take-home message\n\n{take_home}\n\n## 140-character social-media summary\n\n{tweet}\n\n" + "\n\n".join(f"## {h}\n\n{t}" for h,t in body_sections) + "\n\n## References\n\n" + "\n".join(f"{i}. {r}" for i,r in enumerate(refs,1))
    (OUT / "ICM_manuscript_draft.md").write_text(md)

    checklist = f"""# Intensive Care Medicine submission compliance\n\n- Article type: Original paper\n- Main-text word count: {word_count(main_text)} / 3000 maximum\n- Structured abstract: {word_count(abstract_text)} words / 150–250 required\n- Main display items: 5 planned (2 tables + 3 figures) / 5 maximum\n- References: {len(refs)} / 40 maximum\n- Keywords: {len(keywords)} / 4–6 required\n- Take-home message: present, two sentences\n- 140-character social-media summary: {len(tweet)} characters\n- Citations: numbered square-bracket style in manuscript text\n- Headings: <=3 displayed levels\n- Figures: color, blue-toned data graphics\n- Word manuscript: generated as ICM_manuscript_draft.docx\n\n## Items still requiring author input before submission\n\n- Final author list/order and affiliations\n- Corresponding-author email/telephone\n- Penn State IRB protocol number or determination\n- Funding statement\n- Competing-interest statements for all authors\n- Public repository URL and final data/code availability wording\n- Final authorship-contribution statement\n"""
    (OUT / "ICM_submission_checklist.md").write_text(checklist)
    cover = """# Cover letter draft — Intensive Care Medicine\n\nDear Editor-in-Chief,\n\nWe submit the original paper “Day-3 antibiotic de-escalation in suspected ICU infection: a target trial emulation” for consideration in Intensive Care Medicine. The study addresses a common critical-care stewardship decision: whether broad-spectrum antibiotics can be de-escalated at day 3 in patients without a positive clinical culture result available at the decision time.\n\nThe principal contribution is methodological and clinical. In MIMIC-IV, an apparent mortality advantage under limited adjustment progressively disappeared after accounting for near-decision clinical status, improvement trajectories and treatment intensity. This directly demonstrates how recovery-related confounding can distort observational comparisons of de-escalation. A prespecified modified external replication in Penn State clinical data transported the decision/landmark framework but produced a more favorable mortality association under materially different measurement semantics. Across both datasets, lower subsequent antibiotic burden was the most consistent result.\n\nWe believe the manuscript fits ICM because it combines a clinically important ICU stewardship question with a target-trial framework, explicit trajectory adjustment, transparent weighting diagnostics and a frozen external replication. The manuscript has been compressed to the journal’s original-paper format, with technical phenotype, balance, sensitivity and reporting details moved to electronic supplementary material.\n\nThis work has not been published in full elsewhere and is not under consideration by another journal. [Insert ethics, funding, conflicts, and author-approval statements.]\n\nSincerely,\nAlireza Vafaei Sadr\nCorresponding author\n"""
    (OUT / "ICM_cover_letter.md").write_text(cover)
    counts = pd.DataFrame([{ "component":"main_text", "words":word_count(main_text)}, {"component":"abstract", "words":word_count(abstract_text)}, {"component":"references", "words":len(refs)}, {"component":"tweet_characters", "words":len(tweet)}])
    counts.to_csv(OUT / "ICM_wordcounts.csv", index=False)

    make_timeline(); make_progressive(progressive); make_cross_dataset(mortality, secondary)
    build_docx(title, short_title, abstract_sections, keywords, take_home, tweet, body_sections, refs, table1, table2)


if __name__ == "__main__":
    main()
